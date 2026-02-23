import csv
import io
import os
import time
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Set


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".md"}


@dataclass
class StoredDocument:
    file_id: str
    session_id: str
    filename: str
    content_type: str
    text: str
    chunks: List[str]
    created_at: float


class DocumentStore:
    def __init__(self) -> None:
        self.documents: Dict[str, StoredDocument] = {}
        self.session_docs: Dict[str, Set[str]] = {}
        self.max_upload_mb = int(os.getenv("MAX_UPLOAD_MB", "20"))
        self.max_context_chars = int(os.getenv("MAX_DOC_CONTEXT_CHARS", "10000"))
        self.max_chunks = int(os.getenv("MAX_DOC_CHUNKS", "4"))

    def _validate_file(self, filename: str, size_bytes: int) -> None:
        ext = Path(filename).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise ValueError("unsupported_file_type")
        if size_bytes > self.max_upload_mb * 1024 * 1024:
            raise ValueError("file_too_large")

    def _chunk_text(self, text: str, chunk_chars: int = 3000, overlap: int = 300) -> List[str]:
        compact = "\n".join([line.rstrip() for line in text.splitlines() if line.strip()])
        if not compact:
            return []
        chunks: List[str] = []
        start = 0
        n = len(compact)
        while start < n:
            end = min(n, start + chunk_chars)
            chunks.append(compact[start:end])
            if end >= n:
                break
            start = max(0, end - overlap)
        return chunks

    def _extract_text_pdf(self, content: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()

    def _extract_text_docx(self, content: bytes) -> str:
        from docx import Document

        doc = Document(io.BytesIO(content))
        return "\n".join([p.text for p in doc.paragraphs]).strip()

    def _extract_text_txt_like(self, content: bytes) -> str:
        return content.decode("utf-8", errors="ignore").strip()

    def _extract_text_csv(self, content: bytes) -> str:
        raw = content.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(raw))
        rows = []
        for row in reader:
            rows.append(" | ".join([c.strip() for c in row]))
        return "\n".join(rows).strip()

    def extract_text(self, filename: str, content: bytes) -> str:
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return self._extract_text_pdf(content)
        if ext == ".docx":
            return self._extract_text_docx(content)
        if ext in {".txt", ".md"}:
            return self._extract_text_txt_like(content)
        if ext == ".csv":
            return self._extract_text_csv(content)
        raise ValueError("unsupported_file_type")

    def add_document(self, session_id: str, filename: str, content_type: str, content: bytes) -> StoredDocument:
        self._validate_file(filename, len(content))
        text = self.extract_text(filename, content)
        if not text:
            raise ValueError("empty_document")

        file_id = uuid.uuid4().hex
        chunks = self._chunk_text(text)
        doc = StoredDocument(
            file_id=file_id,
            session_id=session_id,
            filename=filename,
            content_type=content_type,
            text=text,
            chunks=chunks,
            created_at=time.time(),
        )
        self.documents[file_id] = doc
        self.session_docs.setdefault(session_id, set()).add(file_id)
        return doc

    def remove_document(self, session_id: str, file_id: str) -> bool:
        doc = self.documents.get(file_id)
        if not doc or doc.session_id != session_id:
            return False
        self.documents.pop(file_id, None)
        if session_id in self.session_docs:
            self.session_docs[session_id].discard(file_id)
        return True

    def list_session_documents(self, session_id: str) -> List[StoredDocument]:
        ids = list(self.session_docs.get(session_id, set()))
        docs = [self.documents[file_id] for file_id in ids if file_id in self.documents]
        return sorted(docs, key=lambda d: d.created_at)

    def _score_chunk(self, chunk: str, query: str) -> int:
        if not query.strip():
            return 0
        query_terms = {t.strip().lower() for t in query.split() if len(t.strip()) > 2}
        lower = chunk.lower()
        return sum(1 for term in query_terms if term in lower)

    def build_context(self, session_id: str, file_ids: Optional[List[str]], query: str) -> str:
        selected: List[StoredDocument] = []
        if file_ids:
            for file_id in file_ids:
                doc = self.documents.get(file_id)
                if doc and doc.session_id == session_id:
                    selected.append(doc)
        else:
            selected = self.list_session_documents(session_id)

        if not selected:
            return ""

        scored: List[tuple[int, str, str]] = []
        for doc in selected:
            for chunk in doc.chunks:
                score = self._score_chunk(chunk, query)
                scored.append((score, doc.filename, chunk))

        scored.sort(key=lambda x: x[0], reverse=True)

        context_parts: List[str] = []
        consumed = 0
        taken = 0
        for score, filename, chunk in scored:
            if taken >= self.max_chunks:
                break
            if consumed + len(chunk) > self.max_context_chars:
                continue
            context_parts.append(f"[Source: {filename}]\n{chunk}")
            consumed += len(chunk)
            taken += 1

        return "\n\n".join(context_parts)


document_store = DocumentStore()
